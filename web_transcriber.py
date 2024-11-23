import streamlit as st
import vosk
import json
import wave
import os
import tempfile
from datetime import datetime
import ffmpeg
from docx import Document
import shutil
import urllib.request
import zipfile

# Aumenta il limite di upload
st.set_page_config(
    page_title="Video Transcriber",
    page_icon="🎥",
    layout="wide",
    menu_items={
        'Get Help': None,
        'Report a bug': None,
        'About': None
    }
)

# Configura il server per file grandi
if not hasattr(st, 'already_started_server'):
    st.already_started_server = True
    st.server.max_upload_size = 5120

@st.cache_resource
def download_model():
    """Scarica e prepara il modello Vosk"""
    model_path = "models/model-it"
    if not os.path.exists(model_path):
        os.makedirs("models", exist_ok=True)
        model_url = "https://alphacephei.com/vosk/models/vosk-model-small-it-0.22.zip"
        zip_path = "model.zip"
        
        with st.spinner("Downloading model... This may take a while..."):
            urllib.request.urlretrieve(model_url, zip_path)
            
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall("models")
            
            if os.path.exists("models/vosk-model-small-it-0.22"):
                os.rename("models/vosk-model-small-it-0.22", model_path)
            
            if os.path.exists(zip_path):
                os.remove(zip_path)
    
    return model_path

def ensure_dirs():
    """Assicura che le directory necessarie esistano"""
    dirs = ['temp', 'output']
    for dir in dirs:
        if not os.path.exists(dir):
            os.makedirs(dir, exist_ok=True)

def extract_audio(video_path, output_path):
    """Estrae l'audio dal video usando ffmpeg"""
    try:
        stream = ffmpeg.input(video_path)
        stream = ffmpeg.output(stream, output_path, 
                             acodec='pcm_s16le', 
                             ac=1, 
                             ar='16k',
                             loglevel='quiet')
        ffmpeg.run(stream, overwrite_output=True)
        return True
    except ffmpeg.Error as e:
        st.error(f"Errore nell'estrazione dell'audio: {str(e)}")
        return False

def process_large_file(uploaded_file, chunk_size=2*1024*1024):
    """Processa file grandi in chunks"""
    total_size = uploaded_file.size
    chunks = []
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as tmp_file:
        with st.progress(0) as progress_bar:
            bytes_read = 0
            while True:
                chunk = uploaded_file.read(chunk_size)
                if not chunk:
                    break
                tmp_file.write(chunk)
                bytes_read += len(chunk)
                progress_bar.progress(min(bytes_read / total_size, 1.0))
        
        return tmp_file.name

def create_transcriber_app():
    ensure_dirs()
    
    st.title("🎥 Video Speech-to-Text")
    st.write("Converti facilmente l'audio dei tuoi video in testo")

    with st.sidebar:
        st.header("Informazioni")
        st.write("""
        ### Come usare:
        1. Seleziona il file video
        2. Attendi il caricamento
        3. Clicca su 'Inizia Trascrizione'
        4. Scarica il documento Word
        
        ### Formati supportati:
        - MP4
        - AVI
        - MKV
        
        ### Note:
        - Supporta video fino a 5GB
        - Non chiudere la finestra durante la trascrizione
        - La trascrizione potrebbe richiedere alcuni minuti
        """)

    uploaded_file = st.file_uploader("Seleziona un video", type=['mp4', 'avi', 'mkv'])
    
    if uploaded_file is not None:
        file_details = {
            "Nome File": uploaded_file.name,
            "Tipo File": uploaded_file.type,
            "Dimensione": f"{uploaded_file.size / (1024*1024):.2f} MB"
        }
        st.write("### Dettagli del file:")
        for key, value in file_details.items():
            st.write(f"**{key}:** {value}")

        if st.button("🎯 Inizia Trascrizione", use_container_width=True):
            try:
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                # Processa il file grande
                status_text.text("⏳ Caricamento del video...")
                temp_video = process_large_file(uploaded_file)
                
                status_text.text("⏳ Estraendo l'audio dal video...")
                
                with tempfile.TemporaryDirectory() as temp_dir:
                    # Estrai audio
                    temp_audio = os.path.join(temp_dir, "audio.wav")
                    if not extract_audio(temp_video, temp_audio):
                        st.error("Errore nell'estrazione dell'audio")
                        return
                    
                    # Prepara documento
                    doc = Document()
                    doc.add_heading('Trascrizione Video', 0)
                    doc.add_paragraph(f'Data: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
                    doc.add_paragraph(f'File originale: {uploaded_file.name}')
                    doc.add_paragraph('')
                    
                    # Carica modello Vosk
                    model_path = download_model()
                    model = vosk.Model(model_path)
                    rec = vosk.KaldiRecognizer(model, 16000)
                    
                    # Trascrivi audio
                    status_text.text("🎯 Trascrizione in corso...")
                    
                    with wave.open(temp_audio, "rb") as wf:
                        total_frames = wf.getnframes()
                        frames_processed = 0
                        
                        while True:
                            data = wf.readframes(4000)
                            if len(data) == 0:
                                break
                            
                            if rec.AcceptWaveform(data):
                                result = json.loads(rec.Result())
                                if result["text"]:
                                    doc.add_paragraph(result["text"])
                            
                            frames_processed += 4000
                            progress = min(frames_processed / total_frames, 1.0)
                            progress_bar.progress(progress)
                    
                    # Aggiungi risultato finale
                    final_result = json.loads(rec.FinalResult())
                    if final_result["text"]:
                        doc.add_paragraph(final_result["text"])
                    
                    # Salva documento
                    output_file = os.path.join("output", f'trascrizione_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
                    doc.save(output_file)
                    
                    # Offri download
                    with open(output_file, 'rb') as f:
                        st.download_button(
                            label="📥 Scarica Trascrizione",
                            data=f,
                            file_name=os.path.basename(output_file),
                            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        )
                    
                    status_text.text("✅ Trascrizione completata!")
                    
                    # Pulizia
                    if os.path.exists(output_file):
                        os.remove(output_file)
                    if os.path.exists(temp_video):
                        os.remove(temp_video)
                
            except Exception as e:
                st.error(f"❌ Errore durante l'elaborazione: {str(e)}")

if __name__ == "__main__":
    create_transcriber_app()